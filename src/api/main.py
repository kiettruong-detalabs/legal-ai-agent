"""
Legal AI Agent API
- Full-text search Vietnamese law database
- Claude OAuth for AI processing
- Multi-tenant API key authentication
- User authentication and management
"""
from fastapi import FastAPI, HTTPException, Depends, Header, Request, Query, UploadFile, File
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import StreamingResponse, FileResponse
from fastapi.staticfiles import StaticFiles
import pathlib
from pydantic import BaseModel, Field
from typing import Optional, List, AsyncGenerator
import psycopg2
from psycopg2.extras import RealDictCursor
import httpx
import json
import hashlib
import time
import re as re_module
import os
from contextlib import contextmanager
from collections import defaultdict
import jwt
from dotenv import load_dotenv

# ============================================
# Rate Limiter (in-memory)
# ============================================
rate_limits: dict = defaultdict(list)

def check_rate_limit(company_id: str, max_requests: int = 20, window_seconds: int = 60):
    """Check if company has exceeded rate limit"""
    now = time.time()
    rate_limits[company_id] = [t for t in rate_limits[company_id] if now - t < window_seconds]
    if len(rate_limits[company_id]) >= max_requests:
        raise HTTPException(status_code=429, detail=f"Rate limit exceeded. Max {max_requests} requests per minute.")
    rate_limits[company_id].append(now)

# ============================================
# Search Cache (in-memory, TTL 1 hour)
# ============================================
search_cache: dict = {}
CACHE_TTL = 3600

# Load environment variables from .env file
load_dotenv()

JWT_SECRET = os.getenv("SUPABASE_JWT_SECRET", "your-super-secret-jwt-key-change-in-production")

# Import new routes
from .routes import auth, company, keys, usage, chats, documents, admin, contracts, templates
# from .middleware.logging import PlatformLoggingMiddleware  # disabled for deploy

# Import agent (initialized after DB functions are defined)
from ..agents import legal_agent

app = FastAPI(
    title="Legal AI Agent API",
    description="AI-powered Vietnamese Legal Assistant - Tư vấn pháp luật, soạn thảo văn bản, rà soát hợp đồng",
    version="2.0.0"
)

# Logging middleware disabled for production deploy
# app.add_middleware(
#     PlatformLoggingMiddleware,
#     exclude_paths=["/health", "/docs", "/openapi.json", "/redoc", "/static", "/"]
# )

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

@app.middleware("http")
async def log_requests(request: Request, call_next):
    start = time.time()
    response = await call_next(request)
    duration = time.time() - start
    if duration > 5:  # Log slow requests
        print(f"SLOW REQUEST: {request.method} {request.url.path} took {duration:.2f}s")
    return response

# Include new routers
app.include_router(auth.router)
app.include_router(company.router)
app.include_router(keys.router)
app.include_router(usage.router)
app.include_router(chats.router)
app.include_router(documents.router)
app.include_router(admin.router)
app.include_router(contracts.router)
app.include_router(templates.router)

# Startup event - seed templates
@app.on_event("startup")
async def startup_event():
    """Seed default templates on startup"""
    templates.seed_default_templates()

# Static files
static_dir = pathlib.Path(__file__).parent.parent.parent / "static"
if static_dir.exists():
    app.mount("/static", StaticFiles(directory=str(static_dir)), name="static")

@app.get("/", include_in_schema=False)
async def landing_page():
    html_file = static_dir / "index.html"
    if html_file.exists():
        return FileResponse(str(html_file))
    return {"name": "Legal AI Agent API", "version": "1.0.0"}

# ============================================
# Database
# ============================================

DB_CONFIG = {
    "host": os.getenv("SUPABASE_DB_HOST", "db.chiokotzjtjwfodryfdt.supabase.co"),
    "port": int(os.getenv("SUPABASE_DB_PORT", "5432")),
    "dbname": "postgres",
    "user": "postgres",
    "password": os.getenv("SUPABASE_DB_PASSWORD", "Hl120804@.,?"),
    "sslmode": "require"
}

@contextmanager
def get_db():
    conn = psycopg2.connect(**DB_CONFIG)
    try:
        yield conn
    finally:
        conn.close()

# ============================================
# Auth
# ============================================

async def verify_api_key(
    x_api_key: Optional[str] = Header(None, alias="X-API-Key"),
    authorization: Optional[str] = Header(None)
):
    """Verify API key OR Bearer token and return company info"""
    
    # Try Bearer token first (from dashboard login)
    if not x_api_key and authorization and authorization.startswith("Bearer "):
        try:
            token = authorization.split(" ", 1)[1]
            payload = jwt.decode(token, JWT_SECRET, algorithms=["HS256"])
            user_id = payload.get("user_id")
            if user_id:
                with get_db() as conn:
                    cur = conn.cursor(cursor_factory=RealDictCursor)
                    cur.execute("""
                        SELECT u.id as user_id, u.company_id, u.role,
                               c.name as company_name, c.plan, c.monthly_quota, c.used_quota
                        FROM users u
                        JOIN companies c ON c.id = u.company_id
                        WHERE u.id = %s
                    """, (user_id,))
                    user = cur.fetchone()
                    if user:
                        if user["used_quota"] >= user["monthly_quota"]:
                            raise HTTPException(status_code=429, detail="Monthly quota exceeded")
                        return {**dict(user), "permissions": ["read","ask","review","draft"], "rate_limit": 60}
        except Exception:
            pass
    
    if not x_api_key:
        raise HTTPException(status_code=401, detail="API key or Bearer token required")
    
    key_hash = hashlib.sha256(x_api_key.encode()).hexdigest()
    key_prefix = x_api_key[:8]
    
    with get_db() as conn:
        cur = conn.cursor(cursor_factory=RealDictCursor)
        cur.execute("""
            SELECT ak.id, ak.company_id, ak.permissions, ak.rate_limit,
                   c.name as company_name, c.plan, c.monthly_quota, c.used_quota
            FROM api_keys ak
            JOIN companies c ON c.id = ak.company_id
            WHERE ak.key_prefix = %s AND ak.key_hash = %s AND ak.is_active = true
        """, (key_prefix, key_hash))
        result = cur.fetchone()
        
        if not result:
            raise HTTPException(status_code=401, detail="Invalid API key")
        
        if result["used_quota"] >= result["monthly_quota"]:
            raise HTTPException(status_code=429, detail="Monthly quota exceeded")
        
        # Update last_used
        cur.execute("UPDATE api_keys SET last_used_at = now() WHERE id = %s", (result["id"],))
        conn.commit()
        
        return dict(result)

# ============================================
# Models
# ============================================

class FileContext(BaseModel):
    filename: str
    content: str
    file_type: str

class LegalQuery(BaseModel):
    question: str = Field(..., min_length=5, max_length=2000, description="Câu hỏi pháp luật")
    domains: Optional[List[str]] = Field(None, description="Lĩnh vực: lao_dong, doanh_nghiep, dan_su, thue, dat_dai...")
    max_sources: int = Field(10, ge=1, le=30, description="Số nguồn tham chiếu tối đa")
    stream: bool = Field(False, description="Stream response")
    session_id: Optional[str] = Field(None, description="Chat session ID for multi-turn conversation")
    file_context: Optional[FileContext] = Field(None, description="File content uploaded in chat")

class ContractReview(BaseModel):
    contract_text: str = Field(..., min_length=50, max_length=100000, description="Nội dung hợp đồng cần rà soát")
    contract_type: Optional[str] = Field(None, description="Loại hợp đồng: hop_dong_lao_dong, hop_dong_thuong_mai...")
    focus_areas: Optional[List[str]] = Field(None, description="Các điểm cần chú ý đặc biệt")

class DocumentDraft(BaseModel):
    doc_type: str = Field(..., description="Loại văn bản: hop_dong_lao_dong, quyet_dinh, cong_van, noi_quy...")
    variables: dict = Field(..., description="Thông tin cần điền vào văn bản")
    instructions: Optional[str] = Field(None, description="Yêu cầu bổ sung")

class LegalResponse(BaseModel):
    answer: str
    citations: List[dict]
    confidence: float
    tokens_used: int
    model: str
    session_id: Optional[str] = None

# ============================================
# Claude OAuth Integration
# ============================================

CLAUDE_OAUTH_TOKEN = os.getenv("CLAUDE_OAUTH_TOKEN", "")
ANTHROPIC_API_KEY = os.getenv("ANTHROPIC_API_KEY", "")
CLAUDE_API_URL = "https://api.anthropic.com/v1/messages"

async def call_claude(system_prompt: str, user_message: str, max_tokens: int = 4096, history: list = None) -> dict:
    """Call Claude via OAuth token or API key, with optional conversation history"""
    api_key = ANTHROPIC_API_KEY
    oauth_token = CLAUDE_OAUTH_TOKEN
    
    headers = {
        "anthropic-version": "2023-06-01",
        "content-type": "application/json"
    }
    
    if oauth_token:
        headers["Authorization"] = f"Bearer {oauth_token}"
        headers["anthropic-beta"] = "oauth-2025-04-20"
    elif api_key:
        headers["x-api-key"] = api_key
    else:
        raise ValueError("No Claude API key or OAuth token configured")
    
    # Build messages with history for multi-turn conversations
    messages = []
    if history:
        for msg in history:
            messages.append({"role": msg["role"], "content": msg["content"]})
    messages.append({"role": "user", "content": user_message})
    
    payload = {
        "model": "claude-sonnet-4-20250514",
        "max_tokens": max_tokens,
        "system": system_prompt,
        "messages": messages
    }
    
    async with httpx.AsyncClient(timeout=120) as client:
        try:
            response = await client.post(CLAUDE_API_URL, headers=headers, json=payload)
            response.raise_for_status()
            data = response.json()
            
            return {
                "content": data["content"][0]["text"],
                "input_tokens": data["usage"]["input_tokens"],
                "output_tokens": data["usage"]["output_tokens"],
                "model": data["model"]
            }
        except httpx.TimeoutException:
            print(f"Claude API timeout")
            raise HTTPException(status_code=504, detail="Hệ thống đang bận, vui lòng thử lại")
        except httpx.HTTPStatusError as e:
            print(f"Claude API error: {e.response.status_code} - {e.response.text[:200]}")
            if e.response.status_code == 429:
                raise HTTPException(status_code=429, detail="Đã vượt giới hạn API, vui lòng chờ")
            elif e.response.status_code == 529:
                raise HTTPException(status_code=503, detail="Hệ thống đang bận, vui lòng thử lại")
            raise HTTPException(status_code=502, detail=f"Lỗi AI engine: {e.response.status_code}")
        except psycopg2.Error as e:
            print(f"Database error in Claude call: {e}")
            raise HTTPException(status_code=500, detail="Lỗi kết nối database")
        except ValueError as e:
            print(f"Claude config error: {e}")
            raise HTTPException(status_code=500, detail=str(e))
        except Exception as e:
            print(f"Claude call error: {e}")
            raise HTTPException(status_code=500, detail="Hệ thống đang bận, vui lòng thử lại")


async def call_claude_stream(system_prompt: str, user_message: str, max_tokens: int = 8192, history: list = None) -> AsyncGenerator[str, None]:
    """Call Claude with streaming via SSE, yielding text deltas"""
    api_key = ANTHROPIC_API_KEY
    oauth_token = CLAUDE_OAUTH_TOKEN

    headers = {
        "anthropic-version": "2023-06-01",
        "content-type": "application/json"
    }

    if oauth_token:
        headers["Authorization"] = f"Bearer {oauth_token}"
        headers["anthropic-beta"] = "oauth-2025-04-20"
    elif api_key:
        headers["x-api-key"] = api_key
    else:
        raise ValueError("No Claude API key or OAuth token configured")

    messages = []
    if history:
        for msg in history:
            messages.append({"role": msg["role"], "content": msg["content"]})
    messages.append({"role": "user", "content": user_message})

    payload = {
        "model": "claude-sonnet-4-20250514",
        "max_tokens": max_tokens,
        "system": system_prompt,
        "messages": messages,
        "stream": True
    }

    async with httpx.AsyncClient(timeout=180) as client:
        async with client.stream("POST", CLAUDE_API_URL, headers=headers, json=payload) as response:
            response.raise_for_status()
            async for line in response.aiter_lines():
                if line.startswith("data: "):
                    data_str = line[6:]
                    if data_str.strip() == "[DONE]":
                        break
                    try:
                        event = json.loads(data_str)
                        event_type = event.get("type", "")
                        if event_type == "content_block_delta":
                            delta = event.get("delta", {})
                            if delta.get("type") == "text_delta":
                                yield delta.get("text", "")
                        elif event_type == "message_stop":
                            break
                        elif event_type == "message_delta":
                            # Contains usage info at the end
                            pass
                    except json.JSONDecodeError:
                        continue


# ============================================
# Context Enrichment - Company Documents & Contracts
# ============================================

def fetch_company_context(company_id: str, question: str, limit: int = 5) -> str:
    """Search company's uploaded documents and contracts for relevant context"""
    context_parts = []
    question_lower = question.lower()

    with get_db() as conn:
        cur = conn.cursor(cursor_factory=RealDictCursor)

        # Search company documents with extracted text
        cur.execute("""
            SELECT name, extracted_text, doc_type, analysis
            FROM documents
            WHERE company_id = %s
              AND extracted_text IS NOT NULL
              AND length(extracted_text) > 50
            ORDER BY created_at DESC
            LIMIT %s
        """, (company_id, limit))

        docs = cur.fetchall()
        for doc in docs:
            text = doc["extracted_text"] or ""
            # Check if document is relevant to the question
            if any(kw in text.lower() for kw in question_lower.split() if len(kw) > 2):
                excerpt = text[:1500]
                context_parts.append(
                    f"📄 TÀI LIỆU CÔNG TY: {doc['name']} (Loại: {doc.get('doc_type', 'N/A')})\n{excerpt}"
                )

        # Search company contracts
        cur.execute("""
            SELECT name, contract_type, extracted_text, parties, start_date, end_date, notes
            FROM contracts
            WHERE company_id = %s
              AND status != 'deleted'
              AND extracted_text IS NOT NULL
              AND length(extracted_text) > 50
            ORDER BY created_at DESC
            LIMIT %s
        """, (company_id, limit))

        contracts = cur.fetchall()
        for contract in contracts:
            text = contract["extracted_text"] or ""
            if any(kw in text.lower() for kw in question_lower.split() if len(kw) > 2):
                excerpt = text[:1500]
                parties_str = ""
                if contract.get("parties"):
                    try:
                        parties = json.loads(contract["parties"]) if isinstance(contract["parties"], str) else contract["parties"]
                        parties_str = f" | Các bên: {', '.join(str(p) for p in parties)}"
                    except:
                        pass
                context_parts.append(
                    f"📋 HỢP ĐỒNG: {contract['name']} (Loại: {contract.get('contract_type', 'N/A')}{parties_str})\n{excerpt}"
                )

    return "\n\n".join(context_parts[:5]) if context_parts else ""


# ============================================
# Law Search
# ============================================

# ============================================
# Vietnamese Diacritics Restoration
# ============================================

NO_ACCENT_MAP = {
    # Common legal phrases
    "thu viec": "thử việc",
    "nghi phep": "nghỉ phép",
    "hop dong lao dong": "hợp đồng lao động",
    "thue tndn": "thuế TNDN",
    "thue thu nhap": "thuế thu nhập",
    "sa thai": "sa thải",
    "luong": "lương",
    "bao hiem": "bảo hiểm",
    "nghi viec": "nghỉ việc",
    "ky luat": "kỷ luật",
    "thai san": "thai sản",
    "tang ca": "tăng ca",
    "lam them gio": "làm thêm giờ",
    "nghi le": "nghỉ lễ",
    "cham dut hop dong": "chấm dứt hợp đồng",
    "boi thuong": "bồi thường",
    "tranh chap": "tranh chấp",
    "thanh lap cong ty": "thành lập công ty",
    "doanh nghiep": "doanh nghiệp",
    "co phan": "cổ phần",
    "thue gia tri gia tang": "thuế giá trị gia tăng",
    "dat dai": "đất đai",
    "quyen su dung dat": "quyền sử dụng đất",
    # Common single words
    "thoi gian": "thời gian",
    "toi da": "tối đa",
    "toi thieu": "tối thiểu",
    "quy dinh": "quy định",
    "noi dung": "nội dung",
    "hinh thuc": "hình thức",
    "hop dong": "hợp đồng",
    "cong ty": "công ty",
    "dieu": "điều",
    "khoan": "khoản",
    "luat": "luật",
    "bo luat": "bộ luật",
    "nghi dinh": "nghị định",
    "thong tu": "thông tư",
    "quyet dinh": "quyết định",
    "muc": "mức",
    "so": "số",
    "nam": "năm",
    "thang": "tháng",
    "ngay": "ngày",
    "gio": "giờ",
    "viec": "việc",
    "nguoi": "người",
    "phep": "phép",
    "thue": "thuế",
    "suat": "suất",
    "tien": "tiền",
}

def has_vietnamese_diacritics(text: str) -> bool:
    """Check if text contains Vietnamese diacritics"""
    import re
    # Vietnamese diacritics pattern
    vietnamese_chars = r'[àáảãạăằắẳẵặâầấẩẫậèéẻẽẹêềếểễệìíỉĩịòóỏõọôồốổỗộơờớởỡợùúủũụưừứửữựỳýỷỹỵđ]'
    return bool(re.search(vietnamese_chars, text.lower()))

def restore_diacritics(query: str) -> str:
    """Restore Vietnamese diacritics from common non-diacritics legal terms"""
    import re
    
    # If query already has diacritics, return as-is
    if has_vietnamese_diacritics(query):
        return query
    
    # Try to match and replace phrases from NO_ACCENT_MAP
    restored = query.lower()
    
    # Sort by length (longest first) to match longer phrases first
    sorted_mappings = sorted(NO_ACCENT_MAP.items(), key=lambda x: len(x[0]), reverse=True)
    
    for no_accent, with_accent in sorted_mappings:
        # Use word boundaries to avoid partial matches
        pattern = r'\b' + re.escape(no_accent) + r'\b'
        restored = re.sub(pattern, with_accent, restored, flags=re.IGNORECASE)
    
    return restored

def extract_search_query(question: str) -> str:
    """Extract key legal terms from Vietnamese question"""
    import re
    
    # Remove Vietnamese question words
    question_words = [
        r'\bbao lâu\b', r'\bbao nhiêu\b', r'\bthế nào\b', r'\bnhư thế nào\b',
        r'\blà gì\b', r'\bcó phải\b', r'\bcó được\b', r'\blà\b', r'\bcó\b',
        r'\bkhông\b', r'\bhay không\b', r'\?', r'\.'
    ]
    
    cleaned = question.lower()
    for pattern in question_words:
        cleaned = re.sub(pattern, ' ', cleaned)
    
    # Remove extra spaces
    cleaned = ' '.join(cleaned.split())
    
    return cleaned.strip()

def expand_synonyms(query: str) -> List[str]:
    """Expand Vietnamese legal term synonyms — returns the EXPANDED TERMS to search for"""
    expansions = []
    query_lower = query.lower()
    
    synonym_map = {
        "tndn": "thu nhập doanh nghiệp",
        "tncn": "thu nhập cá nhân",
        "gtgt": "giá trị gia tăng",
        "vat": "giá trị gia tăng",
        "bhxh": "bảo hiểm xã hội",
        "bhyt": "bảo hiểm y tế",
        "hđlđ": "hợp đồng lao động",
        "nsdlđ": "người sử dụng lao động",
        "nlđ": "người lao động",
        "nghỉ phép": "nghỉ hằng năm",
        "phép năm": "nghỉ hằng năm",
        "sa thải": "kỷ luật sa thải",
        "đuổi việc": "kỷ luật sa thải",
    }
    
    for abbr, full in synonym_map.items():
        if abbr in query_lower:
            expansions.append(full)  # Return just the expanded term
    
    return expansions

def detect_domain(question: str) -> Optional[List[str]]:
    """Auto-detect legal domain from question keywords"""
    question_lower = question.lower()
    
    domain_keywords = {
        "lao_dong": ["lao động", "hợp đồng lao động", "thử việc", "nghỉ phép", "tăng ca", "lương", "sa thải", "bảo hiểm xã hội", "bhxh", "thôi việc", "chấm dứt hợp đồng"],
        "thue": ["thuế", "tndn", "vat", "tncn", "kê khai thuế", "hoàn thuế", "miễn thuế", "giảm thuế", "thuế suất"],
        "doanh_nghiep": ["thành lập công ty", "cổ phần", "doanh nghiệp", "giải thể", "phá sản", "điều lệ", "đại hội cổ đông", "hội đồng quản trị"],
        "dan_su": ["di sản", "thừa kế", "hôn nhân", "ly hôn", "nuôi con", "nhà ở", "quyền sở hữu", "tài sản chung"],
        "dat_dai": ["đất đai", "quyền sử dụng đất", "sổ đỏ", "chuyển nhượng đất", "thuê đất"],
        "hinh_su": ["hình sự", "án tù", "tội phạm", "vi phạm hình sự", "truy tố"],
        "hanh_chinh": ["vi phạm hành chính", "phạt hành chính", "khiếu nại", "tố cáo"]
    }
    
    detected = []
    for domain, keywords in domain_keywords.items():
        for keyword in keywords:
            if keyword in question_lower:
                detected.append(domain)
                break
    
    return detected if detected else None

def search_laws(query: str, domains: Optional[List[str]] = None, limit: int = 10) -> List[dict]:
    """Search Vietnamese law database"""
    with get_db() as conn:
        cur = conn.cursor(cursor_factory=RealDictCursor)
        
        if domains:
            domain_array = "{" + ",".join(domains) + "}"
            cur.execute(
                "SELECT * FROM search_law(%s, %s::legal_domain[], %s)",
                (query, domain_array, limit)
            )
        else:
            cur.execute(
                "SELECT * FROM search_law(%s, NULL, %s)",
                (query, limit)
            )
        
        return [dict(r) for r in cur.fetchall()]

def multi_query_search(question: str, domains: Optional[List[str]] = None, limit: int = 15) -> List[dict]:
    """Smart multi-query search: domain detection + ILIKE phrase + tsvector + diacritics fallback"""
    
    # Auto-detect domain
    if not domains:
        domains = detect_domain(question)
    
    # Check if query needs diacritics restoration
    original_query = question
    restored_query = restore_diacritics(question)
    queries_to_search = [original_query]
    
    # If restoration produced a different query, search with both
    if restored_query.lower() != original_query.lower():
        queries_to_search.append(restored_query)
    
    # Extract clean keywords from the best available query
    keywords = extract_search_query(restored_query if restored_query != original_query else question)
    words = [w for w in keywords.split() if len(w) > 1]
    
    # Build meaningful phrases (skip common words like thời gian, quy định)
    common_prefixes = {"thời", "gian", "quy", "định", "mức", "tối", "đa", "số", "ngày"}
    key_words = [w for w in words if w not in common_prefixes]
    if not key_words:
        key_words = words
    
    # Build search phrases from key words
    phrases = []
    if len(key_words) >= 2:
        phrases.append(" ".join(key_words[:3]))  # Top 3 key words
        phrases.append(" ".join(key_words[:2]))  # Top 2 key words
    elif key_words:
        phrases.append(key_words[0])
    
    # Also try full keyword string
    if len(words) >= 2:
        phrases.append(" ".join(words[:4]))
    
    # Phase 1: ILIKE phrase search with domain filter (most precise)
    phrase_results = []
    with get_db() as conn:
        cur = conn.cursor(cursor_factory=RealDictCursor)
        
        # Search with original phrases
        for phrase in phrases[:3]:
            domain_filter = ""
            params = [f"%{phrase}%"]
            if domains:
                domain_filter = "AND lc.domains && %s::legal_domain[]"
                params.append("{" + ",".join(domains) + "}")
            
            cur.execute(f"""
                SELECT lc.id as chunk_id, lc.law_id, ld.title as law_title, 
                       ld.law_number, lc.article, lc.title as chunk_title,
                       lc.content, lc.domains, 1.0::float as rank
                FROM law_chunks lc
                JOIN law_documents ld ON ld.id = lc.law_id
                WHERE lc.content ILIKE %s {domain_filter}
                ORDER BY 
                    CASE WHEN ld.title LIKE 'Bo Luat%%' OR ld.title LIKE 'Bộ luật%%' THEN 0
                         WHEN ld.title LIKE 'Luat %%' OR ld.title LIKE 'Luật %%' THEN 1
                         WHEN ld.title LIKE 'Nghi dinh%%' OR ld.title LIKE 'Nghị định%%' THEN 2
                         ELSE 3 END,
                    length(lc.content) DESC
                LIMIT {limit}
            """, params)
            phrase_results.extend([dict(r) for r in cur.fetchall()])
        
        # If we have a restored query that's different, search with restored phrases too
        if restored_query.lower() != original_query.lower():
            restored_keywords = extract_search_query(restored_query)
            restored_words = [w for w in restored_keywords.split() if len(w) > 1]
            common_prefixes = {"thời", "gian", "quy", "định", "mức", "tối", "đa", "số", "ngày"}
            restored_key_words = [w for w in restored_words if w not in common_prefixes]
            if not restored_key_words:
                restored_key_words = restored_words
            
            restored_phrases = []
            if len(restored_key_words) >= 2:
                restored_phrases.append(" ".join(restored_key_words[:3]))
                restored_phrases.append(" ".join(restored_key_words[:2]))
            elif restored_key_words:
                restored_phrases.append(restored_key_words[0])
            
            for phrase in restored_phrases[:2]:
                domain_filter = ""
                params = [f"%{phrase}%"]
                if domains:
                    domain_filter = "AND lc.domains && %s::legal_domain[]"
                    params.append("{" + ",".join(domains) + "}")
                
                cur.execute(f"""
                    SELECT lc.id as chunk_id, lc.law_id, ld.title as law_title, 
                           ld.law_number, lc.article, lc.title as chunk_title,
                           lc.content, lc.domains, 1.2::float as rank
                    FROM law_chunks lc
                    JOIN law_documents ld ON ld.id = lc.law_id
                    WHERE lc.content ILIKE %s {domain_filter}
                    ORDER BY 
                        CASE WHEN ld.title LIKE 'Bo Luat%%' OR ld.title LIKE 'Bộ luật%%' THEN 0
                             WHEN ld.title LIKE 'Luat %%' OR ld.title LIKE 'Luật %%' THEN 1
                             WHEN ld.title LIKE 'Nghi dinh%%' OR ld.title LIKE 'Nghị định%%' THEN 2
                             ELSE 3 END,
                        length(lc.content) DESC
                    LIMIT {limit}
                """, params)
                phrase_results.extend([dict(r) for r in cur.fetchall()])
    
    # Phase 1.5: Synonym expansion search (e.g., "tndn" → "thu nhập doanh nghiệp")
    synonyms = expand_synonyms(keywords)
    for syn_term in synonyms[:2]:
        sp = syn_term  # Use the expanded term directly as search phrase
        if True:
            domain_filter = ""
            params = [f"%{sp}%"]
            if domains:
                domain_filter = "AND lc.domains && %s::legal_domain[]"
                params.append("{" + ",".join(domains) + "}")
            with get_db() as conn:
                cur = conn.cursor(cursor_factory=RealDictCursor)
                cur.execute(f"""
                    SELECT lc.id as chunk_id, lc.law_id, ld.title as law_title, 
                           ld.law_number, lc.article, lc.title as chunk_title,
                           lc.content, lc.domains, 1.0::float as rank
                    FROM law_chunks lc
                    JOIN law_documents ld ON ld.id = lc.law_id
                    WHERE lc.content ILIKE %s {domain_filter}
                    ORDER BY CASE WHEN ld.title LIKE 'Bo Luat%%' THEN 0 WHEN ld.title LIKE 'Luat%%' THEN 1 ELSE 2 END
                    LIMIT {limit}
                """, params)
                phrase_results.extend([dict(r) for r in cur.fetchall()])
    
    # Phase 2: tsvector search (broader coverage)
    tsv_results = search_laws(keywords, domains, limit)
    
    # Merge: phrase results first, then tsvector
    seen_ids = set()
    merged = []
    
    # Build title matching keywords from original question
    title_keywords = [w for w in words if len(w) > 2]
    synonym_terms = expand_synonyms(keywords)
    for st in synonym_terms:
        title_keywords.extend(st.split())
    
    for result in phrase_results:
        chunk_id = result.get("chunk_id")
        if chunk_id and chunk_id not in seen_ids:
            seen_ids.add(chunk_id)
            title = result.get("law_title", "").lower()
            base_rank = 15.0
            if any(x in result.get("law_title", "") for x in ["Bo Luat", "Bộ luật"]):
                base_rank = 30.0
            elif any(x in result.get("law_title", "") for x in ["Luat ", "Luật "]):
                base_rank = 25.0
            elif any(x in result.get("law_title", "") for x in ["Nghi dinh", "Nghị định"]):
                base_rank = 20.0
            elif result.get("law_title", "").startswith("Legal Document"):
                base_rank = 8.0
            
            # Boost if law title contains search keywords
            title_match_bonus = 0
            for kw in title_keywords:
                if kw.lower() in title:
                    title_match_bonus += 3.0
            
            result["rank"] = base_rank + title_match_bonus
            merged.append(result)
    
    for result in tsv_results:
        chunk_id = result.get("chunk_id") or result.get("id")
        if chunk_id and chunk_id not in seen_ids:
            seen_ids.add(chunk_id)
            title = result.get("law_title", "")
            base = result.get("rank", 1.0)
            if any(x in title for x in ["Bo Luat", "Bộ luật"]):
                result["rank"] = base + 10.0
            elif any(x in title for x in ["Luat ", "Luật "]):
                result["rank"] = base + 5.0
            elif title.startswith("Legal Document"):
                result["rank"] = max(base - 5.0, 0.1)
            merged.append(result)
    
    merged.sort(key=lambda x: x.get("rank", 0), reverse=True)
    return merged[:limit]


def cached_search(query: str, domains=None, limit=10):
    """Cached wrapper around multi_query_search with 1-hour TTL"""
    domains_key = ",".join(sorted(domains)) if domains else "none"
    cache_key = hashlib.md5(f"{query}:{domains_key}:{limit}".encode()).hexdigest()
    now = time.time()
    if cache_key in search_cache:
        result, timestamp = search_cache[cache_key]
        if now - timestamp < CACHE_TTL:
            return result
    result = multi_query_search(query, domains, limit)
    search_cache[cache_key] = (result, now)
    # Limit cache size
    if len(search_cache) > 1000:
        oldest = min(search_cache, key=lambda k: search_cache[k][1])
        del search_cache[oldest]
    return result


# ============================================
# Initialize Agent with shared functions
# ============================================
legal_agent.init_agent(
    get_db_fn=get_db,
    multi_query_search_fn=multi_query_search,
    search_laws_fn=search_laws,
    detect_domain_fn=detect_domain,
    fetch_company_context_fn=fetch_company_context
)

# ============================================
# API Endpoints
# ============================================

# Root endpoint moved to landing page above

@app.get("/favicon.ico", include_in_schema=False)
@app.get("/favicon.svg", include_in_schema=False)
async def favicon():
    return FileResponse(str(static_dir / "favicon.svg"), media_type="image/svg+xml")

@app.get("/health")
@app.get("/v1/health")
async def health():
    status = {"status": "ok", "version": "2.0.0", "timestamp": time.time(), "ai_engine": "claude-sonnet-4"}
    try:
        with get_db() as conn:
            cur = conn.cursor()
            cur.execute("SELECT count(*) FROM law_documents")
            status["documents"] = cur.fetchone()[0]
            cur.execute("SELECT count(*) FROM law_chunks")
            status["chunks"] = cur.fetchone()[0]
            cur.execute("SELECT count(*) FROM companies")
            status["companies"] = cur.fetchone()[0]
            status["database"] = "connected"
    except Exception:
        status["database"] = "error"
        status["status"] = "degraded"
    return status

@app.post("/v1/chat/upload")
async def chat_upload_file(file: UploadFile = File(...), company: dict = Depends(verify_api_key)):
    """Upload file in chat - extracts text and returns it"""
    from .routes.contracts import extract_file_text
    import tempfile

    # Validate file type
    filename = file.filename or "unknown"
    file_ext = os.path.splitext(filename)[1].lower()
    allowed = {".pdf", ".docx", ".doc", ".txt"}
    if file_ext not in allowed:
        raise HTTPException(status_code=400, detail=f"Loại file không hỗ trợ: {file_ext}. Chỉ chấp nhận: {', '.join(allowed)}")

    content = await file.read()
    if len(content) > 10 * 1024 * 1024:  # 10MB
        raise HTTPException(status_code=400, detail="File quá lớn. Tối đa 10MB.")

    # Extract text
    extracted_text = None
    if file_ext == ".txt":
        extracted_text = content.decode('utf-8', errors='ignore')
    else:
        # Write to temp file for extraction
        with tempfile.NamedTemporaryFile(suffix=file_ext, delete=False) as tmp:
            tmp.write(content)
            tmp_path = tmp.name
        try:
            extracted_text = extract_file_text(tmp_path, file_ext, content)
        finally:
            os.unlink(tmp_path)

    if not extracted_text or len(extracted_text.strip()) < 10:
        raise HTTPException(status_code=422, detail="Không thể trích xuất nội dung từ file. Vui lòng thử file khác.")

    return {
        "filename": filename,
        "content": extracted_text,
        "file_type": file_ext
    }


@app.post("/v1/legal/ask", response_model=LegalResponse)
async def legal_ask(query: LegalQuery, company: dict = Depends(verify_api_key)):
    """Tư vấn pháp luật - Legal Q&A (Agent-based with tool use)"""
    check_rate_limit(str(company["company_id"]))
    
    # Load chat history for multi-turn conversation
    chat_history = []
    session_id = None
    user_id = company.get("user_id")
    
    if query.session_id and user_id:
        try:
            with get_db() as conn:
                cur = conn.cursor(cursor_factory=RealDictCursor)
                cur.execute("""
                    SELECT role, content FROM messages
                    WHERE session_id = %s AND company_id = %s
                    ORDER BY created_at ASC
                    LIMIT 20
                """, (query.session_id, company["company_id"]))
                rows = cur.fetchall()
                for row in rows:
                    chat_history.append({"role": row["role"], "content": row["content"]})
                session_id = query.session_id
        except Exception as e:
            print(f"Error loading chat history: {e}")
    
    # Augment question with file context if provided
    actual_question = query.question
    if query.file_context:
        actual_question = f"""[Người dùng đã upload file: {query.file_context.filename}]

NỘI DUNG FILE:
{query.file_context.content[:30000]}

CÂU HỎI: {query.question}"""

    # Run the agent
    result = await legal_agent.run_agent(
        question=actual_question,
        company_id=str(company["company_id"]),
        user_id=str(user_id) if user_id else None,
        session_id=str(session_id) if session_id else None,
        chat_history=chat_history
    )
    
    citations = result.get("citations", [])
    
    # Save to chat history
    if user_id:
        with get_db() as conn:
            cur = conn.cursor(cursor_factory=RealDictCursor)
            
            if not session_id:
                cur.execute("""
                    SELECT id FROM chat_sessions
                    WHERE user_id = %s AND company_id = %s AND agent_type = 'qa' AND status = 'active'
                    ORDER BY last_message_at DESC NULLS LAST
                    LIMIT 1
                """, (user_id, company["company_id"]))
                
                session = cur.fetchone()
                if session:
                    session_id = session["id"]
                else:
                    cur.execute("""
                        INSERT INTO chat_sessions (company_id, user_id, agent_type, title, status)
                        VALUES (%s, %s, 'qa', %s, 'active')
                        RETURNING id
                    """, (company["company_id"], user_id, f"Q&A - {query.question[:50]}..."))
                    session_id = cur.fetchone()["id"]
            
            cur.execute("""
                INSERT INTO messages (session_id, company_id, role, content, tokens_used, model)
                VALUES (%s, %s, 'user', %s, 0, '')
            """, (session_id, company["company_id"], query.question))
            
            total_tokens = result.get("input_tokens", 0) + result.get("output_tokens", 0)
            cur.execute("""
                INSERT INTO messages (session_id, company_id, role, content, citations, confidence, tokens_used, model)
                VALUES (%s, %s, 'assistant', %s, %s, %s, %s, %s)
            """, (
                session_id,
                company["company_id"],
                result["answer"], 
                json.dumps(citations),
                0.85 if citations else 0.5,
                total_tokens,
                result.get("model", "claude-sonnet-4-20250514")
            ))
            
            cur.execute("""
                UPDATE chat_sessions 
                SET message_count = message_count + 2, last_message_at = now()
                WHERE id = %s
            """, (session_id,))
            
            conn.commit()
    
    # Update usage
    with get_db() as conn:
        cur = conn.cursor()
        cur.execute("UPDATE companies SET used_quota = used_quota + 1 WHERE id = %s", (company["company_id"],))
        cur.execute("""
            INSERT INTO usage_logs (company_id, endpoint, agent_type, input_tokens, output_tokens, status_code)
            VALUES (%s, '/v1/legal/ask', 'agent', %s, %s, 200)
        """, (company["company_id"], result.get("input_tokens", 0), result.get("output_tokens", 0)))
        conn.commit()
    
    return LegalResponse(
        answer=result["answer"],
        citations=citations,
        confidence=0.85 if citations else 0.5,
        tokens_used=result.get("input_tokens", 0) + result.get("output_tokens", 0),
        model=result.get("model", "claude-sonnet-4-20250514"),
        session_id=str(session_id) if session_id else None
    )


@app.post("/v1/legal/ask-stream")
async def legal_ask_stream(query: LegalQuery, company: dict = Depends(verify_api_key)):
    """Tư vấn pháp luật với streaming SSE - Agent-based with tool status events"""
    check_rate_limit(str(company["company_id"]))

    # Load chat history
    chat_history = []
    session_id = None
    user_id = company.get("user_id")

    if query.session_id and user_id:
        try:
            with get_db() as conn:
                cur = conn.cursor(cursor_factory=RealDictCursor)
                cur.execute("""
                    SELECT role, content FROM messages
                    WHERE session_id = %s AND company_id = %s
                    ORDER BY created_at ASC
                    LIMIT 20
                """, (query.session_id, company["company_id"]))
                rows = cur.fetchall()
                for row in rows:
                    chat_history.append({"role": row["role"], "content": row["content"]})
                session_id = query.session_id
        except Exception as e:
            print(f"Error loading chat history: {e}")

    # Augment question with file context if provided
    actual_question = query.question
    if query.file_context:
        actual_question = f"""[Người dùng đã upload file: {query.file_context.filename}]

NỘI DUNG FILE:
{query.file_context.content[:30000]}

CÂU HỎI: {query.question}"""

    company_id_str = str(company["company_id"])

    async def sse_generator():
        """Generate SSE events using agent streaming"""
        full_response = []
        all_citations = []

        try:
            async for event_str in legal_agent.run_agent_stream_final_text(
                question=actual_question,
                company_id=company_id_str,
                user_id=str(user_id) if user_id else None,
                session_id=str(session_id) if session_id else None,
                chat_history=chat_history
            ):
                yield event_str

                # Parse to collect full text and citations for saving
                if event_str.startswith("data: "):
                    try:
                        evt = json.loads(event_str[6:].strip())
                        if evt.get("type") == "delta":
                            full_response.append(evt.get("text", ""))
                        elif evt.get("type") == "citations":
                            all_citations = evt.get("citations", [])
                        elif evt.get("type") == "done":
                            if evt.get("citations"):
                                all_citations = evt["citations"]
                    except:
                        pass

        except Exception as e:
            yield f"data: {json.dumps({'type': 'error', 'message': str(e)}, ensure_ascii=False)}\n\n"
            return

        complete_text = "".join(full_response)

        # Save to chat history
        saved_session_id = session_id
        if user_id and complete_text:
            try:
                with get_db() as conn:
                    cur = conn.cursor(cursor_factory=RealDictCursor)
                    if not saved_session_id:
                        cur.execute("""
                            SELECT id FROM chat_sessions
                            WHERE user_id = %s AND company_id = %s AND agent_type = 'qa' AND status = 'active'
                            ORDER BY last_message_at DESC NULLS LAST LIMIT 1
                        """, (user_id, company["company_id"]))
                        session = cur.fetchone()
                        if session:
                            saved_session_id = session["id"]
                        else:
                            cur.execute("""
                                INSERT INTO chat_sessions (company_id, user_id, agent_type, title, status)
                                VALUES (%s, %s, 'qa', %s, 'active') RETURNING id
                            """, (company["company_id"], user_id, f"Q&A - {query.question[:50]}..."))
                            saved_session_id = cur.fetchone()["id"]

                    cur.execute("""
                        INSERT INTO messages (session_id, company_id, role, content, tokens_used, model)
                        VALUES (%s, %s, 'user', %s, 0, '')
                    """, (saved_session_id, company["company_id"], query.question))

                    cur.execute("""
                        INSERT INTO messages (session_id, company_id, role, content, citations, confidence, tokens_used, model)
                        VALUES (%s, %s, 'assistant', %s, %s, %s, 0, 'claude-sonnet-4-20250514')
                    """, (
                        saved_session_id, company["company_id"],
                        complete_text, json.dumps(all_citations),
                        0.85 if all_citations else 0.5
                    ))

                    cur.execute("""
                        UPDATE chat_sessions SET message_count = message_count + 2, last_message_at = now()
                        WHERE id = %s
                    """, (saved_session_id,))
                    conn.commit()
            except Exception as e:
                print(f"Error saving stream chat history: {e}")

        # Update usage
        try:
            with get_db() as conn:
                cur = conn.cursor()
                cur.execute("UPDATE companies SET used_quota = used_quota + 1 WHERE id = %s",
                            (company["company_id"],))
                cur.execute("""
                    INSERT INTO usage_logs (company_id, endpoint, agent_type, input_tokens, output_tokens, status_code)
                    VALUES (%s, '/v1/legal/ask-stream', 'agent', 0, 0, 200)
                """, (company["company_id"],))
                conn.commit()
        except Exception as e:
            print(f"Error updating usage: {e}")

    return StreamingResponse(
        sse_generator(),
        media_type="text/event-stream",
        headers={
            "Cache-Control": "no-cache",
            "Connection": "keep-alive",
            "X-Accel-Buffering": "no",
        }
    )


@app.get("/v1/legal/search-detailed")
async def search_detailed(
    q: str,
    domains: Optional[str] = None,
    limit: int = Query(20, ge=1, le=50),
    company: dict = Depends(verify_api_key)
):
    """Tìm kiếm luật chi tiết - Detailed Law Search with categorization and highlighting"""
    import time as _time

    start_time = _time.time()

    domain_list = domains.split(",") if domains else None

    # Use cached_search for better results with caching
    results = cached_search(q, domain_list, min(limit, 50))

    elapsed = _time.time() - start_time

    # Categorize results by law type
    categories = {}
    highlighted_results = []
    query_words = [w.lower() for w in q.split() if len(w) > 1]

    for r in results:
        law_title = r.get("law_title", "")
        law_number = r.get("law_number", "")
        content = r.get("content", "")

        # Determine category
        if any(x in law_title for x in ["Bộ luật", "Bo Luat"]):
            cat = "Bộ luật"
        elif any(x in law_title for x in ["Luật ", "Luat "]):
            cat = "Luật"
        elif any(x in law_title for x in ["Nghị định", "Nghi dinh"]):
            cat = "Nghị định"
        elif any(x in law_title for x in ["Thông tư", "Thong tu"]):
            cat = "Thông tư"
        elif any(x in law_title for x in ["Quyết định", "Quyet dinh"]):
            cat = "Quyết định"
        else:
            cat = "Khác"

        if cat not in categories:
            categories[cat] = []

        # Highlight matching text
        highlighted_content = content
        for word in query_words:
            if len(word) > 2:
                pattern = re_module.compile(re_module.escape(word), re_module.IGNORECASE)
                highlighted_content = pattern.sub(f"<mark>{word}</mark>", highlighted_content)

        result_item = {
            "law_title": law_title,
            "law_number": law_number,
            "article": r.get("article"),
            "chunk_title": r.get("chunk_title", ""),
            "content": content[:1000],
            "highlighted_content": highlighted_content[:1000],
            "rank": float(r.get("rank", 0)),
            "category": cat,
            "domains": r.get("domains", [])
        }

        categories[cat].append(result_item)
        highlighted_results.append(result_item)

    # Find related articles (same law, adjacent articles) for top results
    related_articles = []
    if results:
        top_law_ids = list(set(r.get("law_id") for r in results[:3] if r.get("law_id")))
        top_articles = [r.get("article") for r in results[:3] if r.get("article")]

        if top_law_ids:
            try:
                with get_db() as conn:
                    cur = conn.cursor(cursor_factory=RealDictCursor)
                    # Get adjacent articles from the same laws
                    for law_id in top_law_ids[:2]:
                        cur.execute("""
                            SELECT lc.article, lc.title as chunk_title, ld.title as law_title,
                                   ld.law_number, LEFT(lc.content, 300) as content_preview
                            FROM law_chunks lc
                            JOIN law_documents ld ON ld.id = lc.law_id
                            WHERE lc.law_id = %s
                              AND lc.article IS NOT NULL
                            ORDER BY lc.article
                            LIMIT 10
                        """, (law_id,))
                        rows = cur.fetchall()
                        for row in rows:
                            if row["article"] not in top_articles:
                                related_articles.append({
                                    "law_title": row["law_title"],
                                    "law_number": row["law_number"],
                                    "article": row["article"],
                                    "chunk_title": row["chunk_title"],
                                    "content_preview": row["content_preview"]
                                })
            except Exception as e:
                print(f"Error fetching related articles: {e}")

    # Category counts
    category_stats = {cat: len(items) for cat, items in categories.items()}

    return {
        "query": q,
        "stats": {
            "total_results": len(results),
            "search_time_seconds": round(elapsed, 3),
            "categories": category_stats,
            "domains_searched": domain_list
        },
        "results": highlighted_results,
        "categories": {cat: items[:10] for cat, items in categories.items()},
        "related_articles": related_articles[:10]
    }


@app.post("/v1/legal/review")
async def contract_review(review: ContractReview, company: dict = Depends(verify_api_key)):
    """Rà soát hợp đồng - Contract Review"""
    check_rate_limit(str(company["company_id"]))
    
    # Search relevant laws based on contract type
    search_terms = {
        "hop_dong_lao_dong": "hợp đồng lao động quyền nghĩa vụ",
        "hop_dong_thuong_mai": "hợp đồng thương mại mua bán",
        "hop_dong_dich_vu": "hợp đồng dịch vụ thuê khoán",
    }
    search_query = search_terms.get(review.contract_type, "hợp đồng điều khoản")
    sources = cached_search(search_query, None, 15)
    
    context = "\n\n".join([
        f"[{src['law_title']}] {src.get('article', '')}\n{src['content'][:1500]}"
        for src in sources
    ])
    
    system_prompt = """Bạn là luật sư chuyên rà soát hợp đồng theo pháp luật Việt Nam.

Nhiệm vụ: Rà soát hợp đồng và đánh giá theo các tiêu chí:
1. **Tính hợp pháp**: Có điều khoản nào vi phạm pháp luật không?
2. **Tính đầy đủ**: Có thiếu điều khoản bắt buộc nào không?
3. **Rủi ro**: Những điều khoản nào có rủi ro cao cho bên nào?
4. **Đề xuất**: Các sửa đổi cần thiết

Trả về JSON format:
{
    "risk_score": 1-100 (100 = rủi ro cao nhất),
    "issues": [{"type": "violation|missing|risk|suggestion", "severity": "critical|high|medium|low", "clause": "điều khoản liên quan", "description": "mô tả", "legal_basis": "căn cứ pháp lý", "recommendation": "đề xuất sửa"}],
    "summary": "Tóm tắt đánh giá",
    "overall_assessment": "Đánh giá tổng thể"
}"""

    user_message = f"""HỢP ĐỒNG CẦN RÀ SOÁT:
{review.contract_text[:50000]}

PHÁP LUẬT LIÊN QUAN:
{context}

{f"YÊU CẦU ĐẶC BIỆT: {', '.join(review.focus_areas)}" if review.focus_areas else ""}

Hãy rà soát hợp đồng trên và trả về kết quả theo format JSON."""

    result = await call_claude(system_prompt, user_message, max_tokens=8192)
    
    # Update usage
    with get_db() as conn:
        cur = conn.cursor()
        cur.execute("UPDATE companies SET used_quota = used_quota + 1 WHERE id = %s", (company["company_id"],))
        cur.execute("""
            INSERT INTO usage_logs (company_id, endpoint, agent_type, input_tokens, output_tokens, status_code)
            VALUES (%s, '/v1/legal/review', 'review', %s, %s, 200)
        """, (company["company_id"], result["input_tokens"], result["output_tokens"]))
        conn.commit()
    
    # Try to parse JSON from response
    try:
        review_data = json.loads(result["content"])
    except:
        review_data = {"raw_analysis": result["content"]}
    
    return {
        "review": review_data,
        "tokens_used": result["input_tokens"] + result["output_tokens"],
        "model": result["model"]
    }

@app.post("/v1/legal/draft")
async def document_draft(draft: DocumentDraft, company: dict = Depends(verify_api_key)):
    """Soạn thảo văn bản - Document Drafting"""
    check_rate_limit(str(company["company_id"]))
    
    # Search for templates and relevant laws
    sources = cached_search(draft.doc_type.replace("_", " "), None, 10)
    
    context = "\n\n".join([
        f"[{src['law_title']}] {src.get('article', '')}\n{src['content'][:1500]}"
        for src in sources
    ])
    
    system_prompt = """Bạn là chuyên gia soạn thảo văn bản pháp lý Việt Nam.

Nhiệm vụ: Soạn thảo văn bản hoàn chỉnh, đúng format, đúng pháp luật.

Quy tắc:
1. Sử dụng đúng format văn bản hành chính Việt Nam
2. Tuân thủ quy định tại Nghị định 30/2020/NĐ-CP về công tác văn thư
3. Điền đầy đủ thông tin từ biến số được cung cấp
4. Các điều khoản phải tuân thủ pháp luật hiện hành
5. Ghi rõ căn cứ pháp lý"""

    variables_str = json.dumps(draft.variables, ensure_ascii=False, indent=2)
    
    user_message = f"""LOẠI VĂN BẢN: {draft.doc_type}

THÔNG TIN:
{variables_str}

{f"YÊU CẦU BỔ SUNG: {draft.instructions}" if draft.instructions else ""}

PHÁP LUẬT LIÊN QUAN:
{context}

Hãy soạn thảo văn bản hoàn chỉnh."""

    result = await call_claude(system_prompt, user_message, max_tokens=8192)
    
    # Update usage
    with get_db() as conn:
        cur = conn.cursor()
        cur.execute("UPDATE companies SET used_quota = used_quota + 1 WHERE id = %s", (company["company_id"],))
        cur.execute("""
            INSERT INTO usage_logs (company_id, endpoint, agent_type, input_tokens, output_tokens, status_code)
            VALUES (%s, '/v1/legal/draft', 'draft', %s, %s, 200)
        """, (company["company_id"], result["input_tokens"], result["output_tokens"]))
        conn.commit()
    
    return {
        "document": result["content"],
        "doc_type": draft.doc_type,
        "tokens_used": result["input_tokens"] + result["output_tokens"],
        "model": result["model"]
    }

@app.get("/v1/legal/search")
async def search(q: str, domains: Optional[str] = None, limit: int = 10, company: dict = Depends(verify_api_key)):
    """Tìm kiếm luật - Law Search"""
    domain_list = domains.split(",") if domains else None
    results = cached_search(q, domain_list, min(limit, 30))
    
    return {
        "query": q,
        "count": len(results),
        "results": [{
            "law_title": r["law_title"],
            "law_number": r["law_number"],
            "article": r.get("article"),
            "content": r["content"][:500],
            "rank": float(r.get("rank", 0))
        } for r in results]
    }

# ============================================
# Admin endpoints (internal)
# ============================================

@app.post("/admin/company", include_in_schema=False)
async def create_company(name: str, slug: str, plan: str = "trial"):
    """Create a new company (admin only)"""
    import secrets
    
    with get_db() as conn:
        cur = conn.cursor(cursor_factory=RealDictCursor)
        
        # Create company
        cur.execute("""
            INSERT INTO companies (name, slug, plan)
            VALUES (%s, %s, %s::plan_type)
            RETURNING id, name, slug, plan, monthly_quota
        """, (name, slug, plan))
        company = dict(cur.fetchone())
        
        # Generate API key
        api_key = f"lak_{secrets.token_hex(24)}"
        key_hash = hashlib.sha256(api_key.encode()).hexdigest()
        
        cur.execute("""
            INSERT INTO api_keys (company_id, name, key_hash, key_prefix)
            VALUES (%s, %s, %s, %s)
        """, (company["id"], f"{name} - Default Key", key_hash, api_key[:8]))
        
        conn.commit()
        
        return {
            "company": company,
            "api_key": api_key,
            "warning": "Save this API key - it cannot be retrieved later"
        }

# ============================================
# Export to DOCX
# ============================================

class ExportRequest(BaseModel):
    content: str  # Markdown content from AI response
    filename: Optional[str] = "legal-document"

@app.post("/v1/legal/export-docx")
async def export_docx(req: ExportRequest):
    """Convert markdown content to a professional .docx file"""
    import io
    import re
    from docx import Document
    from docx.shared import Pt, Inches, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    
    doc = Document()
    
    # Page margins
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(3)
        section.right_margin = Cm(2)
    
    # Default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(13)
    
    # Heading styles
    for level in range(1, 4):
        h_style = doc.styles[f'Heading {level}']
        h_font = h_style.font
        h_font.name = 'Times New Roman'
        h_font.bold = True
        h_font.color.rgb = RGBColor(0, 0, 0)
        if level == 1:
            h_font.size = Pt(16)
        elif level == 2:
            h_font.size = Pt(14)
        else:
            h_font.size = Pt(13)
    
    lines = req.content.split('\n')
    i = 0
    while i < len(lines):
        line = lines[i].rstrip()
        
        # Skip empty lines
        if not line:
            i += 1
            continue
        
        # Headings
        if line.startswith('### '):
            text = line[4:].strip().strip('*')
            doc.add_heading(text, level=3)
        elif line.startswith('## '):
            text = line[3:].strip().strip('*')
            doc.add_heading(text, level=2)
        elif line.startswith('# '):
            text = line[2:].strip().strip('*')
            doc.add_heading(text, level=1)
        elif line.startswith('---'):
            # Horizontal rule - add thin line
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run('─' * 50)
            run.font.size = Pt(8)
            run.font.color.rgb = RGBColor(150, 150, 150)
        elif line.startswith('- ') or line.startswith('* '):
            # Bullet list
            text = line[2:].strip()
            p = doc.add_paragraph(style='List Bullet')
            _add_formatted_text(p, text)
        elif re.match(r'^\d+[\.\)] ', line):
            # Numbered list
            text = re.sub(r'^\d+[\.\)] ', '', line).strip()
            p = doc.add_paragraph(style='List Number')
            _add_formatted_text(p, text)
        else:
            # Normal paragraph
            p = doc.add_paragraph()
            _add_formatted_text(p, line)
        
        i += 1
    
    # Save to bytes
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    
    safe_name = re.sub(r'[^\w\-]', '_', req.filename)
    
    return StreamingResponse(
        buffer,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={
            "Content-Disposition": f'attachment; filename="{safe_name}.docx"'
        }
    )

def _add_formatted_text(paragraph, text: str):
    """Parse markdown bold/italic and add formatted runs to paragraph"""
    import re
    from docx.shared import Pt
    
    # Split by **bold** and *italic* patterns
    parts = re.split(r'(\*\*.*?\*\*|\*.*?\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
            run.font.name = 'Times New Roman'
            run.font.size = Pt(13)
        elif part.startswith('*') and part.endswith('*'):
            run = paragraph.add_run(part[1:-1])
            run.italic = True
            run.font.name = 'Times New Roman'
            run.font.size = Pt(13)
        elif part.startswith('[') and ']' in part:
            # Placeholder like [TÊN CÔNG TY] - highlight it
            run = paragraph.add_run(part)
            run.bold = True
            run.font.name = 'Times New Roman'
            run.font.size = Pt(13)
            from docx.shared import RGBColor
            run.font.color.rgb = RGBColor(200, 0, 0)  # Red for placeholders
        else:
            # Check for [PLACEHOLDER] within normal text
            sub_parts = re.split(r'(\[.*?\])', part)
            for sp in sub_parts:
                if sp.startswith('[') and sp.endswith(']'):
                    run = paragraph.add_run(sp)
                    run.bold = True
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(13)
                    from docx.shared import RGBColor
                    run.font.color.rgb = RGBColor(200, 0, 0)
                else:
                    run = paragraph.add_run(sp)
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(13)


# ============================================
# Feature: Smart Contract Comparison
# ============================================

class ContractCompareRequest(BaseModel):
    contract_ids: List[str] = Field(..., min_length=2, description="Danh sách ID hợp đồng cần so sánh (tối thiểu 2)")

@app.post("/v1/contracts/compare")
async def compare_contracts(req: ContractCompareRequest, company: dict = Depends(verify_api_key)):
    """Compare 2+ contracts side-by-side with AI analysis"""
    check_rate_limit(str(company["company_id"]))

    if len(req.contract_ids) < 2:
        raise HTTPException(status_code=400, detail="Cần ít nhất 2 hợp đồng để so sánh")
    if len(req.contract_ids) > 5:
        raise HTTPException(status_code=400, detail="Tối đa 5 hợp đồng")

    contracts_data = []
    with get_db() as conn:
        cur = conn.cursor(cursor_factory=RealDictCursor)
        for cid in req.contract_ids:
            cur.execute("""
                SELECT id, name, contract_type, extracted_text, parties,
                       start_date, end_date, value, status
                FROM contracts
                WHERE id::text = %s AND company_id = %s AND status != 'deleted'
            """, (cid, company["company_id"]))
            contract = cur.fetchone()
            if not contract:
                raise HTTPException(status_code=404, detail=f"Không tìm thấy hợp đồng: {cid}")
            c = dict(contract)
            for key in ["start_date", "end_date"]:
                if c.get(key):
                    c[key] = str(c[key])
            contracts_data.append(c)

    # Build comparison prompt
    contracts_text = ""
    for i, c in enumerate(contracts_data, 1):
        text = (c.get("extracted_text") or "")[:8000]
        contracts_text += f"\n\n--- HỢP ĐỒNG {i}: {c['name']} (Loại: {c.get('contract_type', 'N/A')}) ---\n{text}"

    system_prompt = """Bạn là luật sư chuyên so sánh hợp đồng. Phân tích và so sánh các hợp đồng được cung cấp.

Trả về JSON format:
{
    "summary": "Tóm tắt so sánh tổng quan",
    "contracts": [{"id": "...", "name": "...", "strengths": ["..."], "weaknesses": ["..."]}],
    "differences": [{"aspect": "Khía cạnh", "details": [{"contract": "Tên HĐ", "content": "Nội dung"}], "recommendation": "Đề xuất"}],
    "inconsistencies": ["Điểm không nhất quán giữa các hợp đồng"],
    "best_for_company": {"contract_name": "...", "reason": "Lý do"},
    "recommendations": ["Đề xuất chung"]
}"""

    user_message = f"So sánh {len(contracts_data)} hợp đồng sau:{contracts_text}\n\nHãy phân tích chi tiết sự khác biệt, điểm mạnh/yếu của từng hợp đồng, và đề xuất hợp đồng nào có lợi hơn cho công ty."

    result = await call_claude(system_prompt, user_message, max_tokens=8192)

    # Update usage
    with get_db() as conn:
        cur = conn.cursor()
        cur.execute("UPDATE companies SET used_quota = used_quota + 1 WHERE id = %s", (company["company_id"],))
        cur.execute("""
            INSERT INTO usage_logs (company_id, endpoint, agent_type, input_tokens, output_tokens, status_code)
            VALUES (%s, '/v1/contracts/compare', 'compare', %s, %s, 200)
        """, (company["company_id"], result["input_tokens"], result["output_tokens"]))
        conn.commit()

    try:
        comparison = json.loads(result["content"])
    except:
        comparison = {"raw_analysis": result["content"]}

    return {
        "comparison": comparison,
        "contracts": [{
            "id": str(c["id"]),
            "name": c["name"],
            "contract_type": c.get("contract_type"),
            "start_date": c.get("start_date"),
            "end_date": c.get("end_date"),
            "value": c.get("value")
        } for c in contracts_data],
        "tokens_used": result["input_tokens"] + result["output_tokens"],
        "model": result["model"]
    }


# ============================================
# Feature: Contract Timeline & Deadline Tracking
# ============================================

@app.get("/v1/contracts/timeline")
async def contract_timeline(company: dict = Depends(verify_api_key)):
    """Get all contract deadlines, expiry dates as a timeline"""
    from datetime import datetime, timedelta

    with get_db() as conn:
        cur = conn.cursor(cursor_factory=RealDictCursor)
        cur.execute("""
            SELECT id, name, contract_type, start_date, end_date, value, status, parties
            FROM contracts
            WHERE company_id = %s AND status != 'deleted'
              AND (start_date IS NOT NULL OR end_date IS NOT NULL)
            ORDER BY
                CASE WHEN end_date IS NOT NULL THEN end_date ELSE start_date END ASC
        """, (company["company_id"],))
        contracts = cur.fetchall()

    now = datetime.now().date()
    timeline_items = []
    overdue_count = 0
    expiring_30 = 0
    expiring_60 = 0
    expiring_90 = 0

    for c in contracts:
        item = dict(c)
        for key in ["start_date", "end_date"]:
            if item.get(key):
                item[key] = str(item[key])
        if item.get("parties"):
            try:
                if isinstance(item["parties"], str):
                    item["parties"] = json.loads(item["parties"])
            except:
                pass

        end_date = c.get("end_date")
        if end_date:
            days_remaining = (end_date - now).days
            item["days_remaining"] = days_remaining
            if days_remaining < 0:
                item["timeline_status"] = "overdue"
                overdue_count += 1
            elif days_remaining <= 30:
                item["timeline_status"] = "expiring_soon"
                expiring_30 += 1
            elif days_remaining <= 60:
                item["timeline_status"] = "expiring_60"
                expiring_60 += 1
            elif days_remaining <= 90:
                item["timeline_status"] = "expiring_90"
                expiring_90 += 1
            else:
                item["timeline_status"] = "active"
        else:
            item["days_remaining"] = None
            item["timeline_status"] = "no_end_date"

        item["id"] = str(item["id"])
        timeline_items.append(item)

    return {
        "timeline": timeline_items,
        "summary": {
            "total": len(timeline_items),
            "overdue": overdue_count,
            "expiring_30_days": expiring_30,
            "expiring_60_days": expiring_60,
            "expiring_90_days": expiring_90,
            "today": str(now)
        }
    }


# ============================================
# Feature: Document Annotations
# ============================================

@app.on_event("startup")
async def create_annotations_table():
    """Create annotations table if not exists"""
    try:
        with get_db() as conn:
            cur = conn.cursor()
            cur.execute("""
                CREATE TABLE IF NOT EXISTS document_annotations (
                    id UUID DEFAULT gen_random_uuid() PRIMARY KEY,
                    document_id UUID NOT NULL,
                    company_id UUID NOT NULL,
                    user_id UUID,
                    text_selection TEXT,
                    start_offset INTEGER,
                    end_offset INTEGER,
                    comment TEXT NOT NULL,
                    annotation_type VARCHAR(50) DEFAULT 'comment',
                    is_ai_generated BOOLEAN DEFAULT FALSE,
                    created_at TIMESTAMPTZ DEFAULT now(),
                    updated_at TIMESTAMPTZ DEFAULT now()
                )
            """)
            cur.execute("""
                CREATE INDEX IF NOT EXISTS idx_annotations_document
                ON document_annotations(document_id, company_id)
            """)
            conn.commit()
    except Exception as e:
        print(f"Annotations table creation: {e}")


class AnnotationCreate(BaseModel):
    text_selection: Optional[str] = None
    start_offset: Optional[int] = None
    end_offset: Optional[int] = None
    comment: str = Field(..., min_length=1, max_length=5000)
    annotation_type: str = Field("comment", description="comment, issue, suggestion, highlight")


@app.post("/v1/documents/{doc_id}/annotate")
async def annotate_document(doc_id: str, annotation: AnnotationCreate, company: dict = Depends(verify_api_key)):
    """Add annotation/comment to a specific part of document"""
    user_id = company.get("user_id")

    with get_db() as conn:
        cur = conn.cursor(cursor_factory=RealDictCursor)
        # Verify document belongs to company
        cur.execute("""
            SELECT id FROM documents WHERE id::text = %s AND company_id = %s
            UNION ALL
            SELECT id FROM contracts WHERE id::text = %s AND company_id = %s AND status != 'deleted'
        """, (doc_id, company["company_id"], doc_id, company["company_id"]))
        if not cur.fetchone():
            raise HTTPException(status_code=404, detail="Tài liệu không tồn tại")

        cur.execute("""
            INSERT INTO document_annotations
                (document_id, company_id, user_id, text_selection, start_offset, end_offset, comment, annotation_type)
            VALUES (%s, %s, %s, %s, %s, %s, %s, %s)
            RETURNING id, document_id, text_selection, start_offset, end_offset, comment, annotation_type, is_ai_generated, created_at
        """, (
            doc_id, company["company_id"], user_id,
            annotation.text_selection, annotation.start_offset, annotation.end_offset,
            annotation.comment, annotation.annotation_type
        ))
        result = dict(cur.fetchone())
        conn.commit()

    result["id"] = str(result["id"])
    result["document_id"] = str(result["document_id"])
    result["created_at"] = str(result["created_at"])
    return result


@app.get("/v1/documents/{doc_id}/annotations")
async def get_annotations(doc_id: str, company: dict = Depends(verify_api_key)):
    """Get all annotations for a document"""
    with get_db() as conn:
        cur = conn.cursor(cursor_factory=RealDictCursor)
        cur.execute("""
            SELECT a.id, a.document_id, a.user_id, a.text_selection, a.start_offset, a.end_offset,
                   a.comment, a.annotation_type, a.is_ai_generated, a.created_at,
                   u.full_name as author_name
            FROM document_annotations a
            LEFT JOIN users u ON u.id = a.user_id
            WHERE a.document_id::text = %s AND a.company_id = %s
            ORDER BY a.start_offset ASC NULLS LAST, a.created_at ASC
        """, (doc_id, company["company_id"]))
        annotations = cur.fetchall()

    result = []
    for ann in annotations:
        item = dict(ann)
        for key in ["id", "document_id", "user_id"]:
            if item.get(key):
                item[key] = str(item[key])
        if item.get("created_at"):
            item["created_at"] = str(item["created_at"])
        result.append(item)

    return {"annotations": result, "total": len(result)}


@app.delete("/v1/documents/{doc_id}/annotations/{annotation_id}")
async def delete_annotation(doc_id: str, annotation_id: str, company: dict = Depends(verify_api_key)):
    """Delete an annotation"""
    with get_db() as conn:
        cur = conn.cursor()
        cur.execute("""
            DELETE FROM document_annotations
            WHERE id::text = %s AND document_id::text = %s AND company_id = %s
        """, (annotation_id, doc_id, company["company_id"]))
        if cur.rowcount == 0:
            raise HTTPException(status_code=404, detail="Annotation không tồn tại")
        conn.commit()
    return {"message": "Đã xóa annotation"}


# ============================================
# Feature: Notifications Endpoint
# ============================================

@app.get("/v1/notifications")
async def get_notifications(company: dict = Depends(verify_api_key)):
    """Get notifications for the company — expiring contracts, overdue items"""
    from datetime import date as date_type
    notifications = []

    with get_db() as conn:
        cur = conn.cursor(cursor_factory=RealDictCursor)

        # Expiring contracts (within 30 days)
        cur.execute("""
            SELECT id, name, end_date
            FROM contracts
            WHERE company_id = %s AND status != 'deleted'
            AND end_date IS NOT NULL
            AND end_date <= CURRENT_DATE + INTERVAL '30 days'
            AND end_date >= CURRENT_DATE
            ORDER BY end_date ASC
        """, (company["company_id"],))

        for c in cur.fetchall():
            days = (c['end_date'] - date_type.today()).days
            notifications.append({
                "type": "expiring_contract",
                "severity": "warning" if days > 7 else "critical",
                "title": f"HĐ sắp hết hạn: {c['name']}",
                "detail": f"Còn {days} ngày",
                "action": {"type": "view_contract", "id": str(c['id'])}
            })

        # Overdue contracts
        cur.execute("""
            SELECT id, name, end_date
            FROM contracts
            WHERE company_id = %s AND status != 'deleted'
            AND end_date IS NOT NULL AND end_date < CURRENT_DATE
            ORDER BY end_date DESC LIMIT 5
        """, (company["company_id"],))

        for c in cur.fetchall():
            notifications.append({
                "type": "overdue_contract",
                "severity": "critical",
                "title": f"HĐ đã hết hạn: {c['name']}",
                "detail": f"Hết hạn {c['end_date']}",
                "action": {"type": "view_contract", "id": str(c['id'])}
            })

    return {"notifications": notifications, "count": len(notifications)}


if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=8080)
