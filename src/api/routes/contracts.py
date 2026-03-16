"""
Contract Management - Multi-tenant CRUD + AI Review
Upload, manage, and review contracts with Claude AI
"""
from fastapi import APIRouter, HTTPException, Depends, UploadFile, File, Form, Query
from typing import Optional, List, Dict
from pydantic import BaseModel
from datetime import datetime, date, timedelta
import os
import json
import uuid
from pathlib import Path
import psycopg2
from psycopg2.extras import RealDictCursor
from ..middleware.auth import get_current_user, get_db
import httpx

router = APIRouter(prefix="/v1/contracts", tags=["contracts"])

# File upload config
UPLOAD_DIR = Path("/home/admin_1/projects/legal-ai-agent/uploads/contracts")
UPLOAD_DIR.mkdir(parents=True, exist_ok=True)
ALLOWED_EXTENSIONS = {".pdf", ".docx", ".doc", ".txt", ".jpg", ".jpeg", ".png"}


def extract_file_text(file_path: str, file_ext: str, content: bytes = None) -> Optional[str]:
    """Extract text from PDF, DOCX, or TXT files"""
    try:
        if file_ext == ".txt":
            return content.decode('utf-8', errors='ignore') if content else ""
        
        elif file_ext == ".pdf":
            from PyPDF2 import PdfReader
            reader = PdfReader(file_path)
            text_parts = []
            for page in reader.pages:
                text = page.extract_text()
                if text:
                    text_parts.append(text)
            return "\n".join(text_parts) if text_parts else None
        
        elif file_ext in (".docx", ".doc"):
            from docx import Document
            doc = Document(file_path)
            text_parts = []
            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text)
            # Also extract from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                    if row_text:
                        text_parts.append(row_text)
            return "\n".join(text_parts) if text_parts else None
        
        elif file_ext in (".jpg", ".jpeg", ".png"):
            # TODO: OCR with pytesseract if available
            return None
    except Exception as e:
        print(f"Text extraction error ({file_ext}): {e}")
        return None


async def ai_analyze_contract(text: str) -> dict:
    """Use Claude to extract contract metadata from text"""
    CLAUDE_OAUTH_TOKEN = os.getenv("CLAUDE_OAUTH_TOKEN", "")
    
    system_prompt = """Phân tích văn bản hợp đồng và trích xuất thông tin. Trả về JSON thuần túy (không markdown):
{
    "title": "Tên hợp đồng (ví dụ: Hợp đồng lao động giữa Công ty ABC và Nguyễn Văn A)",
    "contract_type": "hop_dong_lao_dong|hop_dong_dich_vu|hop_dong_thue|hop_dong_mua_ban|hop_dong_hop_tac|khac",
    "parties": ["Bên A: Công ty ABC", "Bên B: Nguyễn Văn A"],
    "start_date": "YYYY-MM-DD hoặc null",
    "end_date": "YYYY-MM-DD hoặc null",
    "value": số tiền (số nguyên VNĐ) hoặc null,
    "summary": "Tóm tắt nội dung chính trong 2-3 câu"
}
Chỉ trả về JSON, không thêm text hay markdown."""

    headers = {
        "Authorization": f"Bearer {CLAUDE_OAUTH_TOKEN}",
        "anthropic-beta": "oauth-2025-04-20",
        "anthropic-version": "2023-06-01",
        "content-type": "application/json"
    }
    
    payload = {
        "model": "claude-sonnet-4-20250514",
        "max_tokens": 1024,
        "system": system_prompt,
        "messages": [{"role": "user", "content": f"Phân tích hợp đồng:\n\n{text[:15000]}"}]
    }
    
    try:
        async with httpx.AsyncClient(timeout=60) as client:
            response = await client.post(
                "https://api.anthropic.com/v1/messages",
                headers=headers,
                json=payload
            )
            response.raise_for_status()
            data = response.json()
            content = data["content"][0]["text"]
            # Clean markdown code blocks if any
            content = content.strip()
            if content.startswith("```"):
                content = content.split("\n", 1)[1] if "\n" in content else content
                content = content.rsplit("```", 1)[0] if "```" in content else content
            return json.loads(content.strip())
    except Exception as e:
        print(f"AI contract analysis error: {e}")
        return {}


# ============================================
# Models
# ============================================

class ContractCreate(BaseModel):
    name: str
    contract_type: Optional[str] = None
    parties: Optional[List[Dict]] = []
    start_date: Optional[date] = None
    end_date: Optional[date] = None
    notes: Optional[str] = None
    metadata: Optional[Dict] = {}

class ContractUpdate(BaseModel):
    name: Optional[str] = None
    contract_type: Optional[str] = None
    parties: Optional[List[Dict]] = None
    start_date: Optional[date] = None
    end_date: Optional[date] = None
    status: Optional[str] = None
    notes: Optional[str] = None
    metadata: Optional[Dict] = None

class ContractReview(BaseModel):
    contract_id: Optional[str] = None
    contract_text: Optional[str] = None
    focus_areas: Optional[List[str]] = []

# ============================================
# Helpers
# ============================================

async def call_claude_for_review(contract_text: str, contract_type: str = None) -> dict:
    """Call Claude for contract review"""
    CLAUDE_OAUTH_TOKEN = os.getenv("CLAUDE_OAUTH_TOKEN", "")
    
    system_prompt = """Bạn là luật sư chuyên rà soát hợp đồng theo pháp luật Việt Nam.

Nhiệm vụ: Phân tích hợp đồng và đánh giá theo các tiêu chí:
1. **Tính hợp pháp**: Có điều khoản vi phạm pháp luật không?
2. **Tính đầy đủ**: Thiếu điều khoản bắt buộc nào không?
3. **Rủi ro**: Điều khoản nào có rủi ro cao?
4. **Đề xuất**: Sửa đổi cần thiết

Trả về JSON:
{
    "risk_score": 1-100,
    "issues": [{"type": "violation|missing|risk|suggestion", "severity": "critical|high|medium|low", "clause": "...", "description": "...", "legal_basis": "...", "recommendation": "..."}],
    "summary": "Tóm tắt",
    "overall_assessment": "Đánh giá tổng thể"
}"""

    user_message = f"""HỢP ĐỒNG CẦN RÀ SOÁT:
{contract_text[:30000]}

{f"LOẠI HỢP ĐỒNG: {contract_type}" if contract_type else ""}

Hãy rà soát và trả về JSON như yêu cầu."""

    headers = {
        "Authorization": f"Bearer {CLAUDE_OAUTH_TOKEN}",
        "anthropic-beta": "oauth-2025-04-20",
        "anthropic-version": "2023-06-01",
        "content-type": "application/json"
    }
    
    payload = {
        "model": "claude-sonnet-4-20250514",
        "max_tokens": 8192,
        "system": system_prompt,
        "messages": [{"role": "user", "content": user_message}]
    }
    
    async with httpx.AsyncClient(timeout=120) as client:
        response = await client.post(
            "https://api.anthropic.com/v1/messages",
            headers=headers,
            json=payload
        )
        response.raise_for_status()
        data = response.json()
        
        content = data["content"][0]["text"]
        
        # Try to parse JSON from Claude's response
        try:
            review_result = json.loads(content)
        except:
            # If not valid JSON, wrap it
            review_result = {"raw_analysis": content}
        
        return {
            "review": review_result,
            "tokens": {
                "input": data["usage"]["input_tokens"],
                "output": data["usage"]["output_tokens"]
            }
        }

# ============================================
# Endpoints
# ============================================

@router.post("/analyze")
async def analyze_uploaded_file(
    file: UploadFile = File(...),
    current_user: Dict = Depends(get_current_user)
):
    """Upload a file and get AI-extracted metadata (name, type, parties, dates, value)"""
    file_ext = Path(file.filename).suffix.lower()
    if file_ext not in ALLOWED_EXTENSIONS:
        raise HTTPException(400, f"File type not allowed. Allowed: {', '.join(ALLOWED_EXTENSIONS)}")
    
    content = await file.read()
    
    # Save temp file for extraction
    temp_id = str(uuid.uuid4())
    temp_path = str(UPLOAD_DIR / f"temp_{temp_id}{file_ext}")
    with open(temp_path, "wb") as f:
        f.write(content)
    
    # Extract text
    extracted_text = extract_file_text(temp_path, file_ext, content)
    
    # Clean up temp file
    try:
        os.remove(temp_path)
    except:
        pass
    
    if not extracted_text or len(extracted_text.strip()) < 20:
        return {
            "success": False,
            "message": "Không thể đọc nội dung file. Vui lòng nhập thông tin thủ công.",
            "extracted_text": extracted_text,
            "metadata": {}
        }
    
    # AI analyze
    metadata = await ai_analyze_contract(extracted_text)
    
    return {
        "success": True,
        "extracted_text": extracted_text[:5000],
        "metadata": metadata
    }


@router.post("")
async def create_contract(
    name: str = Form(""),
    contract_type: Optional[str] = Form(None),
    parties: Optional[str] = Form("[]"),
    start_date: Optional[str] = Form(None),
    end_date: Optional[str] = Form(None),
    notes: Optional[str] = Form(None),
    value: Optional[str] = Form("0"),
    file: Optional[UploadFile] = File(None),
    current_user: Dict = Depends(get_current_user)
):
    """Create/upload contract with optional AI auto-fill"""
    
    # Parse parties JSON
    try:
        parties_list = json.loads(parties) if parties else []
    except:
        parties_list = []
    
    # Parse dates
    start_date_obj = datetime.strptime(start_date, "%Y-%m-%d").date() if start_date else None
    end_date_obj = datetime.strptime(end_date, "%Y-%m-%d").date() if end_date else None
    
    # Handle file upload
    file_path = None
    file_type = None
    extracted_text = None
    
    if file:
        # Validate file extension
        file_ext = Path(file.filename).suffix.lower()
        if file_ext not in ALLOWED_EXTENSIONS:
            raise HTTPException(
                status_code=400,
                detail=f"File type not allowed. Allowed: {', '.join(ALLOWED_EXTENSIONS)}"
            )
        
        # Save file
        file_id = str(uuid.uuid4())
        file_name = f"{file_id}{file_ext}"
        file_path = str(UPLOAD_DIR / file_name)
        
        with open(file_path, "wb") as f:
            content = await file.read()
            f.write(content)
        
        file_type = file_ext
        
        # Extract text from uploaded file
        extracted_text = extract_file_text(file_path, file_ext, content)
    
    # Insert into database
    with get_db() as conn:
        cur = conn.cursor(cursor_factory=RealDictCursor)
        
        cur.execute("""
            INSERT INTO contracts (
                company_id, uploaded_by, name, contract_type, parties,
                start_date, end_date, file_path, file_type, extracted_text, notes
            )
            VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s)
            RETURNING *
        """, (
            current_user["company_id"],
            current_user["id"],
            name,
            contract_type,
            json.dumps(parties_list),
            start_date_obj,
            end_date_obj,
            file_path,
            file_type,
            extracted_text,
            notes
        ))
        
        contract = dict(cur.fetchone())
        conn.commit()
        
        return contract

@router.get("")
async def list_contracts(
    contract_type: Optional[str] = None,
    status: Optional[str] = None,
    start_date: Optional[str] = None,
    end_date: Optional[str] = None,
    limit: int = Query(50, le=200),
    offset: int = 0,
    current_user: Dict = Depends(get_current_user)
):
    """List contracts for company"""
    with get_db() as conn:
        cur = conn.cursor(cursor_factory=RealDictCursor)
        
        query = """
            SELECT 
                c.*,
                u.full_name as uploaded_by_name,
                u.email as uploaded_by_email
            FROM contracts c
            LEFT JOIN users u ON u.id = c.uploaded_by
            WHERE c.company_id = %s
        """
        params = [current_user["company_id"]]
        
        if contract_type:
            query += " AND c.contract_type = %s"
            params.append(contract_type)
        
        if status:
            query += " AND c.status = %s"
            params.append(status)
        
        if start_date:
            query += " AND c.start_date >= %s"
            params.append(start_date)
        
        if end_date:
            query += " AND c.end_date <= %s"
            params.append(end_date)
        
        query += " ORDER BY c.created_at DESC LIMIT %s OFFSET %s"
        params.extend([limit, offset])
        
        cur.execute(query, params)
        contracts = [dict(r) for r in cur.fetchall()]
        
        return {"contracts": contracts}

@router.get("/expiring")
async def get_expiring_contracts(
    days: int = Query(30, ge=1, le=365),
    current_user: Dict = Depends(get_current_user)
):
    """Get contracts expiring within N days"""
    with get_db() as conn:
        cur = conn.cursor(cursor_factory=RealDictCursor)
        
        cur.execute("""
            SELECT 
                c.*,
                u.full_name as uploaded_by_name,
                (c.end_date - CURRENT_DATE) as days_until_expiry
            FROM contracts c
            LEFT JOIN users u ON u.id = c.uploaded_by
            WHERE c.company_id = %s
              AND c.end_date IS NOT NULL
              AND c.end_date > CURRENT_DATE
              AND c.end_date <= CURRENT_DATE + %s * INTERVAL '1 day'
              AND c.status = 'active'
            ORDER BY c.end_date ASC
        """, (current_user["company_id"], days))
        
        contracts = [dict(r) for r in cur.fetchall()]
        
        return {"expiring_contracts": contracts, "days": days}

@router.get("/{contract_id}")
async def get_contract(contract_id: str, current_user: Dict = Depends(get_current_user)):
    """Get contract detail"""
    with get_db() as conn:
        cur = conn.cursor(cursor_factory=RealDictCursor)
        
        cur.execute("""
            SELECT 
                c.*,
                u.full_name as uploaded_by_name,
                u.email as uploaded_by_email
            FROM contracts c
            LEFT JOIN users u ON u.id = c.uploaded_by
            WHERE c.id = %s AND c.company_id = %s
        """, (contract_id, current_user["company_id"]))
        
        contract = cur.fetchone()
        if not contract:
            raise HTTPException(status_code=404, detail="Contract not found")
        
        return dict(contract)

@router.put("/{contract_id}")
async def update_contract(
    contract_id: str,
    update: ContractUpdate,
    current_user: Dict = Depends(get_current_user)
):
    """Update contract metadata"""
    with get_db() as conn:
        cur = conn.cursor(cursor_factory=RealDictCursor)
        
        # Build update query
        updates = []
        params = []
        
        if update.name is not None:
            updates.append("name = %s")
            params.append(update.name)
        
        if update.contract_type is not None:
            updates.append("contract_type = %s")
            params.append(update.contract_type)
        
        if update.parties is not None:
            updates.append("parties = %s::jsonb")
            params.append(json.dumps(update.parties))
        
        if update.start_date is not None:
            updates.append("start_date = %s")
            params.append(update.start_date)
        
        if update.end_date is not None:
            updates.append("end_date = %s")
            params.append(update.end_date)
        
        if update.status is not None:
            updates.append("status = %s")
            params.append(update.status)
        
        if update.notes is not None:
            updates.append("notes = %s")
            params.append(update.notes)
        
        if update.metadata is not None:
            updates.append("metadata = %s::jsonb")
            params.append(json.dumps(update.metadata))
        
        if not updates:
            raise HTTPException(status_code=400, detail="No updates provided")
        
        updates.append("updated_at = now()")
        
        query = f"""
            UPDATE contracts 
            SET {', '.join(updates)}
            WHERE id = %s AND company_id = %s
            RETURNING *
        """
        params.extend([contract_id, current_user["company_id"]])
        
        cur.execute(query, params)
        updated = cur.fetchone()
        
        if not updated:
            raise HTTPException(status_code=404, detail="Contract not found")
        
        conn.commit()
        return dict(updated)

@router.delete("/{contract_id}")
async def delete_contract(
    contract_id: str,
    current_user: Dict = Depends(get_current_user)
):
    """Soft delete contract"""
    with get_db() as conn:
        cur = conn.cursor(cursor_factory=RealDictCursor)
        
        cur.execute("""
            UPDATE contracts
            SET status = 'deleted', updated_at = now()
            WHERE id = %s AND company_id = %s
            RETURNING id
        """, (contract_id, current_user["company_id"]))
        
        deleted = cur.fetchone()
        if not deleted:
            raise HTTPException(status_code=404, detail="Contract not found")
        
        conn.commit()
        return {"message": "Contract deleted", "id": deleted["id"]}

@router.post("/{contract_id}/review")
async def review_contract(
    contract_id: str,
    current_user: Dict = Depends(get_current_user)
):
    """Send contract for AI review"""
    with get_db() as conn:
        cur = conn.cursor(cursor_factory=RealDictCursor)
        
        # Get contract
        cur.execute("""
            SELECT * FROM contracts
            WHERE id = %s AND company_id = %s
        """, (contract_id, current_user["company_id"]))
        
        contract = cur.fetchone()
        if not contract:
            raise HTTPException(status_code=404, detail="Contract not found")
        
        contract = dict(contract)
        
        # Get text (either extracted or from file)
        text = contract.get("extracted_text")
        
        if not text:
            raise HTTPException(
                status_code=400,
                detail="No text available for review. Upload a text file or extract text from PDF/DOCX first."
            )
        
        # Call Claude for review
        result = await call_claude_for_review(text, contract.get("contract_type"))
        
        # Save review result
        cur.execute("""
            UPDATE contracts
            SET review_result = %s::jsonb, updated_at = now()
            WHERE id = %s
            RETURNING *
        """, (json.dumps(result["review"]), contract_id))
        
        updated = dict(cur.fetchone())
        conn.commit()
        
        # Update usage
        cur.execute("""
            UPDATE companies 
            SET used_quota = used_quota + 1 
            WHERE id = %s
        """, (current_user["company_id"],))
        
        cur.execute("""
            INSERT INTO usage_logs (
                company_id, user_id, endpoint, agent_type,
                input_tokens, output_tokens, status_code
            )
            VALUES (%s, %s, '/v1/contracts/review', 'review', %s, %s, 200)
        """, (
            current_user["company_id"],
            current_user["id"],
            result["tokens"]["input"],
            result["tokens"]["output"]
        ))
        
        conn.commit()
        
        return {
            "contract": updated,
            "review": result["review"],
            "tokens_used": result["tokens"]["input"] + result["tokens"]["output"]
        }

@router.post("/review-text")
async def review_contract_text(
    review: ContractReview,
    current_user: Dict = Depends(get_current_user)
):
    """Review contract from direct text input (no upload)"""
    if not review.contract_text:
        raise HTTPException(status_code=400, detail="contract_text is required")
    
    result = await call_claude_for_review(review.contract_text)
    
    # Update usage
    with get_db() as conn:
        cur = conn.cursor()
        
        cur.execute("""
            UPDATE companies 
            SET used_quota = used_quota + 1 
            WHERE id = %s
        """, (current_user["company_id"],))
        
        cur.execute("""
            INSERT INTO usage_logs (
                company_id, user_id, endpoint, agent_type,
                input_tokens, output_tokens, status_code
            )
            VALUES (%s, %s, '/v1/contracts/review-text', 'review', %s, %s, 200)
        """, (
            current_user["company_id"],
            current_user["id"],
            result["tokens"]["input"],
            result["tokens"]["output"]
        ))
        
        conn.commit()
    
    return {
        "review": result["review"],
        "tokens_used": result["tokens"]["input"] + result["tokens"]["output"]
    }
